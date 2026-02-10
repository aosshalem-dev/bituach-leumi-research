#!/usr/bin/env python3
"""
Convert Hebrew markdown research paper to Word (.docx) with proper RTL formatting.
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_rtl(paragraph):
    """Set paragraph to RTL (Right-to-Left) direction."""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def parse_markdown_line(line):
    """Parse a single line for inline formatting."""
    # Store the original line for processing
    segments = []
    current_pos = 0

    # Find all bold (**text**) and italic (*text*) markers
    # Pattern: **bold** or *italic*
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*)'

    for match in re.finditer(pattern, line):
        # Add text before the match
        if match.start() > current_pos:
            segments.append(('normal', line[current_pos:match.start()]))

        # Add the formatted text
        if match.group(2):  # Bold
            segments.append(('bold', match.group(2)))
        elif match.group(3):  # Italic
            segments.append(('italic', match.group(3)))

        current_pos = match.end()

    # Add remaining text
    if current_pos < len(line):
        segments.append(('normal', line[current_pos:]))

    # If no formatting found, return the whole line as normal
    if not segments:
        segments.append(('normal', line))

    return segments

def add_formatted_paragraph(doc, text, style='Normal'):
    """Add a paragraph with RTL and inline formatting."""
    paragraph = doc.add_paragraph(style=style)
    set_rtl(paragraph)

    segments = parse_markdown_line(text)
    for format_type, content in segments:
        run = paragraph.add_run(content)
        if format_type == 'bold':
            run.bold = True
        elif format_type == 'italic':
            run.italic = True

    return paragraph

def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx."""
    table_lines = []
    idx = start_idx

    while idx < len(lines) and '|' in lines[idx]:
        table_lines.append(lines[idx])
        idx += 1

    if len(table_lines) < 2:
        return None, start_idx

    # Parse header
    header_line = table_lines[0]
    headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]

    # Skip separator line
    rows = []
    for line in table_lines[2:]:
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            rows.append(cells)

    return {'headers': headers, 'rows': rows}, idx

def add_table_to_doc(doc, table_data):
    """Add a table to the document."""
    if not table_data:
        return

    headers = table_data['headers']
    rows = table_data['rows']

    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Add headers
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        paragraph = cell.paragraphs[0]
        set_rtl(paragraph)
        paragraph.text = header
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.runs[0]
        run.bold = True

    # Add rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(row_cells):
                cell = row_cells[col_idx]
                paragraph = cell.paragraphs[0]
                set_rtl(paragraph)
                paragraph.text = cell_data
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def convert_markdown_to_docx(input_file, output_file):
    """Convert markdown file to docx with RTL formatting."""

    # Read the markdown file
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set up styles
    styles = doc.styles

    # Configure Normal style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'David'
    normal_font.size = Pt(12)

    # Configure heading styles
    for i in range(1, 4):
        heading_style = styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = 'David'
        if i == 1:
            heading_font.size = Pt(18)
            heading_font.bold = True
        elif i == 2:
            heading_font.size = Pt(16)
            heading_font.bold = True
        else:
            heading_font.size = Pt(14)
            heading_font.bold = True

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add title page
    title = doc.add_paragraph()
    set_rtl(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('מיהו מבוטח בביטוח הלאומי?')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.name = 'David'

    doc.add_page_break()

    # Split content into lines
    lines = content.split('\n')

    # Collect footnotes
    footnotes = {}
    footnote_pattern = r'^\[\^(\d+)\]:\s*(.+)$'

    # First pass: collect footnotes
    for line in lines:
        footnote_match = re.match(footnote_pattern, line.strip())
        if footnote_match:
            footnote_num = footnote_match.group(1)
            footnote_text = footnote_match.group(2)
            footnotes[footnote_num] = footnote_text

    # Second pass: process content
    i = 0
    in_list = False
    list_items = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Skip footnote definitions
        if re.match(footnote_pattern, line.strip()):
            i += 1
            continue

        # Skip separator lines
        if line.strip() == '---':
            doc.add_paragraph()
            i += 1
            continue

        # Handle headings
        if line.startswith('###'):
            text = line.replace('###', '').strip()
            add_formatted_paragraph(doc, text, 'Heading 3')
            i += 1
            continue
        elif line.startswith('##'):
            text = line.replace('##', '').strip()
            add_formatted_paragraph(doc, text, 'Heading 2')
            i += 1
            continue
        elif line.startswith('#'):
            text = line.replace('#', '').strip()
            add_formatted_paragraph(doc, text, 'Heading 1')
            i += 1
            continue

        # Handle blockquotes
        if line.startswith('>'):
            text = line.replace('>', '').strip()
            paragraph = add_formatted_paragraph(doc, text, 'Normal')
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.right_indent = Inches(0.5)
            i += 1
            continue

        # Handle tables
        if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_data, next_idx = parse_table(lines, i)
            if table_data:
                add_table_to_doc(doc, table_data)
                i = next_idx
                continue

        # Handle numbered lists
        numbered_list_match = re.match(r'^(\d+)\.\s+(.+)$', line)
        if numbered_list_match:
            text = numbered_list_match.group(2)
            # Replace inline footnote references
            text = re.sub(r'\[\^(\d+)\]', r'[\1]', text)
            paragraph = add_formatted_paragraph(doc, text, 'List Number')
            i += 1
            continue

        # Handle bullet points
        if line.strip().startswith('-') or line.strip().startswith('*'):
            text = re.sub(r'^[\-\*]\s+', '', line.strip())
            # Replace inline footnote references
            text = re.sub(r'\[\^(\d+)\]', r'[\1]', text)
            paragraph = add_formatted_paragraph(doc, text, 'List Bullet')
            i += 1
            continue

        # Handle normal paragraphs
        text = line.strip()
        # Replace inline footnote references
        text = re.sub(r'\[\^(\d+)\]', r'[\1]', text)
        add_formatted_paragraph(doc, text, 'Normal')
        i += 1

    # Add footnotes section if there are any
    if footnotes:
        doc.add_page_break()

        # Add footnotes heading
        footnotes_heading = doc.add_paragraph()
        set_rtl(footnotes_heading)
        footnotes_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_run = footnotes_heading.add_run('הערות שוליים')
        heading_run.font.size = Pt(16)
        heading_run.font.bold = True
        heading_run.font.name = 'David'

        doc.add_paragraph()

        # Add footnotes in order
        for num in sorted(footnotes.keys(), key=int):
            footnote_text = f"[{num}] {footnotes[num]}"
            paragraph = add_formatted_paragraph(doc, footnote_text, 'Normal')
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)

    # Save document
    doc.save(output_file)
    print(f"Document saved to: {output_file}")

if __name__ == '__main__':
    input_file = '/Users/zvishalem/Downloads/bituach_leumi_research/reports/full_paper.md'
    output_file = '/Users/zvishalem/Downloads/bituach_leumi_research/reports/full_paper.docx'

    convert_markdown_to_docx(input_file, output_file)
